"""
Generador del .docx de email copys para revisar con Ari.

Lee EMAIL_COPYS.md y produce docs/email-review/Email-Copys.docx.
El body de cada email se renderiza como párrafo natural (sin los
`>` de markdown), indentado, en gris suave — leyéndose como
mail de verdad y no como source code.

Run:
  python3 scripts/build-email-copys-docx.py
"""

from docx import Document
from docx.shared import Pt, RGBColor
import re
from pathlib import Path

SOURCE = Path(__file__).parent.parent / "EMAIL_COPYS.md"
TARGET = Path(__file__).parent.parent / "docs" / "email-review" / "Email-Copys.docx"


def is_blockquote(line: str) -> bool:
    """Treat any line whose first non-space char is '>' as blockquote.
    The .md indents some quotes by 2 spaces."""
    return line.lstrip().startswith(">")


def strip_quote(line: str) -> str:
    """Remove leading whitespace + '> ' / '>'."""
    return line.lstrip()[1:].lstrip()


def flush_quote(doc, buffer):
    """Render buffered quote lines as indented body paragraphs.

    Cada línea del body es su propio párrafo (no las unimos para no
    perder la cadencia natural del email: saludo, body, closing,
    botón). Sin el `>` prefix. Italic + gris suave para diferenciar
    del metadata "Cuándo / Subject / Body:".
    """
    if not buffer:
        return
    for line in buffer:
        if line.strip() == "":
            # Empty quote line → small gap between body paragraphs.
            doc.add_paragraph("")
            continue
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(24)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(strip_markdown(line))
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        run.italic = True
    buffer.clear()


def main():
    md = SOURCE.read_text(encoding="utf-8")
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    quote_buffer: list[str] = []

    for raw_line in md.split("\n"):
        line = raw_line.rstrip()

        # Blockquote (body) accumulation
        if is_blockquote(line):
            quote_buffer.append(strip_quote(line))
            continue
        else:
            flush_quote(doc, quote_buffer)

        # Empty line — paragraph break
        if not line.strip():
            doc.add_paragraph("")
            continue

        # Horizontal rules
        if line.strip() == "---":
            doc.add_paragraph("─" * 40)
            continue

        # Headings
        if line.startswith("### "):
            doc.add_heading(strip_markdown(line[4:]), level=3)
            continue
        if line.startswith("## "):
            doc.add_heading(strip_markdown(line[3:]), level=2)
            continue
        if line.startswith("# "):
            doc.add_heading(strip_markdown(line[2:]), level=1)
            continue

        # Lists
        m = re.match(r"^(\s*)- (.*)$", line)
        if m:
            text = strip_markdown(m.group(2))
            doc.add_paragraph(text, style="List Bullet")
            continue
        m = re.match(r"^(\s*)(\d+)\. (.*)$", line)
        if m:
            text = strip_markdown(m.group(3))
            doc.add_paragraph(text, style="List Number")
            continue

        # Tables (rare here) as plain paragraphs
        if line.startswith("|"):
            doc.add_paragraph(strip_markdown(line))
            continue

        # Plain paragraph with bold runs preserved
        write_inline(doc, line)

    # Flush any trailing quote
    flush_quote(doc, quote_buffer)

    TARGET.parent.mkdir(parents=True, exist_ok=True)
    doc.save(TARGET)
    print(f"Wrote {TARGET}")


def strip_markdown(s: str) -> str:
    s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
    s = re.sub(r"\*(.+?)\*", r"\1", s)
    s = re.sub(r"`(.+?)`", r"\1", s)
    s = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", s)
    return s


def write_inline(doc, line: str):
    p = doc.add_paragraph()
    parts = re.split(r"(\*\*.+?\*\*)", line)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(strip_markdown(part))


if __name__ == "__main__":
    main()
